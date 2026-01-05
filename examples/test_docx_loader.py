"""
Script demo để test DOCX Loader với nhiều loại file DOCX.

Test cases:
1. DOCX text thuần
2. DOCX có formatting (bold, italic, headings)
3. DOCX có bảng
4. DOCX có ảnh embedded (OCR)
5. DOCX có lists và numbering
6. DOCX đa ngôn ngữ
"""

import sys
import os
from pathlib import Path

# Thêm đường dẫn đến thư mục gốc
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from core.loaders.docx_loader import DOCXLoader, load_docx


def create_test_docx_files():
    """Tạo các file DOCX test với python-docx."""
    print("\n" + "=" * 70)
    print("TẠO FILE DOCX TEST")
    print("=" * 70)
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 1. DOCX với text đơn giản
        print("\n📝 Đang tạo test_simple.docx...")
        doc1 = Document()
        doc1.add_heading('Document Test Đơn Giản', 0)
        doc1.add_paragraph('Đây là đoạn văn bản đơn giản bằng tiếng Việt.')
        doc1.add_paragraph('Paragraph thứ hai với nội dung khác.')
        doc1.add_paragraph('Paragraph thứ ba có thêm thông tin.')
        doc1.save('test_simple.docx')
        print("   ✅ Đã tạo test_simple.docx")
        
        # 2. DOCX với formatting
        print("\n📝 Đang tạo test_formatted.docx...")
        doc2 = Document()
        doc2.add_heading('Document Có Formatting', 0)
        doc2.add_heading('Heading Level 1', level=1)
        
        p1 = doc2.add_paragraph('Đây là văn bản có ')
        p1.add_run('chữ đậm').bold = True
        p1.add_run(' và ')
        p1.add_run('chữ nghiêng').italic = True
        p1.add_run('.')
        
        doc2.add_heading('Heading Level 2', level=2)
        doc2.add_paragraph('Paragraph bình thường tiếp theo.')
        doc2.save('test_formatted.docx')
        print("   ✅ Đã tạo test_formatted.docx")
        
        # 3. DOCX với bảng
        print("\n📝 Đang tạo test_table.docx...")
        doc3 = Document()
        doc3.add_heading('Document Có Bảng', 0)
        doc3.add_paragraph('Dưới đây là một bảng dữ liệu:')
        
        table = doc3.add_table(rows=4, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'STT'
        header_cells[1].text = 'Tên'
        header_cells[2].text = 'Điểm'
        
        # Data rows
        data = [
            ('1', 'Nguyễn Văn A', '9.0'),
            ('2', 'Trần Thị B', '8.5'),
            ('3', 'Lê Văn C', '9.5')
        ]
        
        for i, (stt, ten, diem) in enumerate(data, 1):
            cells = table.rows[i].cells
            cells[0].text = stt
            cells[1].text = ten
            cells[2].text = diem
        
        doc3.save('test_table.docx')
        print("   ✅ Đã tạo test_table.docx")
        
        # 4. DOCX với lists
        print("\n📝 Đang tạo test_lists.docx...")
        doc4 = Document()
        doc4.add_heading('Document Có Lists', 0)
        
        doc4.add_paragraph('Danh sách không số:')
        doc4.add_paragraph('Item 1', style='List Bullet')
        doc4.add_paragraph('Item 2', style='List Bullet')
        doc4.add_paragraph('Item 3', style='List Bullet')
        
        doc4.add_paragraph('Danh sách có số:')
        doc4.add_paragraph('Bước 1: Chuẩn bị', style='List Number')
        doc4.add_paragraph('Bước 2: Thực hiện', style='List Number')
        doc4.add_paragraph('Bước 3: Hoàn thành', style='List Number')
        
        doc4.save('test_lists.docx')
        print("   ✅ Đã tạo test_lists.docx")
        
        # 5. DOCX đa ngôn ngữ
        print("\n📝 Đang tạo test_multilang.docx...")
        doc5 = Document()
        doc5.add_heading('Multilingual Document', 0)
        doc5.add_paragraph('English: Hello, World!')
        doc5.add_paragraph('Tiếng Việt: Xin chào thế giới!')
        doc5.add_paragraph('日本語: こんにちは世界！')
        doc5.add_paragraph('한국어: 안녕하세요 세계!')
        doc5.add_paragraph('中文: 你好世界！')
        doc5.save('test_multilang.docx')
        print("   ✅ Đã tạo test_multilang.docx")
        
        print("\n✅ Hoàn thành tạo các file test!")
        return True
        
    except ImportError:
        print("\n❌ Lỗi: Cần cài đặt python-docx")
        print("   Chạy: pip install python-docx")
        return False
    except Exception as e:
        print(f"\n❌ Lỗi khi tạo file: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_loader_initialization():
    """Test khởi tạo DOCX Loader với các config khác nhau."""
    print("\n" + "=" * 70)
    print("TEST 1: KHỞI TẠO DOCX LOADER")
    print("=" * 70)
    
    # Config 1: Default (full features)
    loader1 = DOCXLoader()
    print("\n✅ Config 1 - Default (Full Features):")
    print(f"   Image extraction: {loader1.enable_image_extraction}")
    print(f"   Table extraction: {loader1.enable_table_extraction}")
    print(f"   Text cleaning: {loader1.enable_text_cleaning}")
    
    # Config 2: Text-only (fast mode)
    loader2 = DOCXLoader(
        enable_image_extraction=False,
        enable_table_extraction=False,
        enable_text_cleaning=False
    )
    print("\n✅ Config 2 - Text-Only Mode (Fast):")
    print(f"   Image extraction: {loader2.enable_image_extraction}")
    print(f"   Table extraction: {loader2.enable_table_extraction}")
    
    # Config 3: OCR-focused
    loader3 = DOCXLoader(
        enable_image_extraction=True,
        min_image_confidence=70.0,
        min_image_words=10
    )
    print("\n✅ Config 3 - OCR-Focused:")
    print(f"   Min confidence: {loader3.min_image_confidence}%")
    print(f"   Min words: {loader3.min_image_words}")


def test_basic_docx_loading(docx_path):
    """Test load DOCX cơ bản."""
    if not os.path.exists(docx_path):
        print(f"\n⚠️  File không tồn tại: {docx_path}")
        return None
    
    print("\n" + "=" * 70)
    print("TEST 2: LOAD DOCX CƠ BẢN")
    print("=" * 70)
    print(f"\n📂 File: {docx_path}")
    
    loader = DOCXLoader()
    docs = loader.load_docx(docx_path)
    
    if docs:
        print(f"\n✅ Load thành công!")
        print(f"   Số documents: {len(docs)}")
        
        # Thông tin document
        doc = docs[0]
        metadata = doc['metadata']
        
        print(f"\n📄 Metadata:")
        print(f"   File name: {metadata.get('file_name', 'N/A')}")
        print(f"   File type: {metadata.get('file_type', 'N/A')}")
        print(f"   Encoding: {metadata.get('encoding', 'N/A')}")
        print(f"   Paragraphs: {metadata.get('num_paragraphs', 'N/A')}")
        print(f"   Words: {metadata.get('num_words', 'N/A')}")
        print(f"   Has tables: {metadata.get('has_tables', False)}")
        print(f"   Has images: {metadata.get('has_images', False)}")
        
        print(f"\n📝 Preview nội dung (300 ký tự đầu):")
        print("─" * 70)
        print(doc['text'][:300] + "...")
        print("─" * 70)
        
        return docs
    else:
        print("\n❌ Không load được file!")
        return None


def test_table_extraction(docx_path):
    """Test trích xuất bảng từ DOCX."""
    if not os.path.exists(docx_path):
        print(f"\n⚠️  File không tồn tại: {docx_path}")
        return
    
    print("\n" + "=" * 70)
    print("TEST 3: TRÍCH XUẤT BẢNG")
    print("=" * 70)
    print(f"\n📂 File: {docx_path}")
    
    loader = DOCXLoader(enable_table_extraction=True)
    docs = loader.load_docx(docx_path)
    
    if docs:
        doc = docs[0]
        metadata = doc['metadata']
        
        num_tables = metadata.get('num_tables', 0)
        print(f"\n✅ Tìm thấy {num_tables} bảng")
        
        if num_tables > 0:
            print("\n📊 Nội dung bảng đã được extract và thêm vào text:")
            print("─" * 70)
            # Tìm phần text có bảng
            text = doc['text']
            if 'STT' in text or 'Tên' in text:
                start = text.find('STT')
                if start != -1:
                    print(text[start:start+200] + "...")
            print("─" * 70)


def test_image_extraction(docx_path):
    """Test trích xuất và OCR ảnh trong DOCX."""
    if not os.path.exists(docx_path):
        print(f"\n⚠️  File không tồn tại: {docx_path}")
        return
    
    print("\n" + "=" * 70)
    print("TEST 4: TRÍCH XUẤT ẢNH VÀ OCR")
    print("=" * 70)
    print(f"\n📂 File: {docx_path}")
    
    loader = DOCXLoader(
        enable_image_extraction=True,
        min_image_confidence=60.0,
        min_image_words=5
    )
    docs = loader.load_docx(docx_path)
    
    if docs:
        doc = docs[0]
        metadata = doc['metadata']
        
        num_images = metadata.get('num_images', 0)
        print(f"\n✅ Tìm thấy {num_images} ảnh")
        
        if num_images > 0:
            image_texts = metadata.get('image_text', [])
            print(f"\n📷 Text từ ảnh:")
            for i, img_text in enumerate(image_texts[:3], 1):  # Show 3 ảnh đầu
                print(f"\n   Ảnh {i}:")
                print(f"      {img_text[:100]}...")
        else:
            print("\n   (File này không có ảnh với text)")


def test_formatting_preservation(docx_path):
    """Test bảo toàn formatting."""
    if not os.path.exists(docx_path):
        print(f"\n⚠️  File không tồn tại: {docx_path}")
        return
    
    print("\n" + "=" * 70)
    print("TEST 5: BẢO TOÀN FORMATTING")
    print("=" * 70)
    print(f"\n📂 File: {docx_path}")
    
    loader = DOCXLoader()
    docs = loader.load_docx(docx_path)
    
    if docs:
        doc = docs[0]
        text = doc['text']
        
        print("\n📝 Text đã extract:")
        print("─" * 70)
        print(text)
        print("─" * 70)
        
        print("\n✅ Formatting notes:")
        print("   - Headings được giữ nguyên")
        print("   - Paragraphs được phân tách rõ ràng")
        print("   - Lists được format với bullets/numbers")


def test_multilingual(docx_path):
    """Test DOCX đa ngôn ngữ."""
    if not os.path.exists(docx_path):
        print(f"\n⚠️  File không tồn tại: {docx_path}")
        return
    
    print("\n" + "=" * 70)
    print("TEST 6: DOCX ĐA NGÔN NGỮ")
    print("=" * 70)
    print(f"\n📂 File: {docx_path}")
    
    loader = DOCXLoader()
    docs = loader.load_docx(docx_path)
    
    if docs:
        doc = docs[0]
        metadata = doc['metadata']
        
        print(f"\n✅ Load thành công!")
        print(f"   Encoding: {metadata.get('encoding', 'N/A')}")
        print(f"   Words: {metadata.get('num_words', 'N/A')}")
        
        print(f"\n🌍 Nội dung đa ngôn ngữ:")
        print("─" * 70)
        print(doc['text'])
        print("─" * 70)


def test_text_cleaning():
    """Test text cleaning."""
    print("\n" + "=" * 70)
    print("TEST 7: TEXT CLEANING")
    print("=" * 70)
    
    print("\n✅ Text cleaning sẽ:")
    print("   - Gộp nhiều khoảng trắng thành 1")
    print("   - Xóa dòng trống dư thừa")
    print("   - Xóa URLs và emails (nếu enable)")
    print("   - Chuẩn hóa encoding")
    print("   - Normalize Unicode characters")


def demo_with_created_files():
    """Demo với các file đã tạo."""
    print("\n" + "=" * 70)
    print("DEMO VỚI FILE DOCX ĐÃ TẠO")
    print("=" * 70)
    
    test_files = [
        ("test_simple.docx", "Text đơn giản"),
        ("test_formatted.docx", "Có formatting"),
        ("test_table.docx", "Có bảng"),
        ("test_lists.docx", "Có lists"),
        ("test_multilang.docx", "Đa ngôn ngữ")
    ]
    
    for filename, description in test_files:
        if os.path.exists(filename):
            print(f"\n{'=' * 70}")
            print(f"Testing: {filename} ({description})")
            print('=' * 70)
            test_basic_docx_loading(filename)


def cleanup_test_files():
    """Xóa các file test."""
    test_files = [
        "test_simple.docx",
        "test_formatted.docx",
        "test_table.docx",
        "test_lists.docx",
        "test_multilang.docx"
    ]
    
    print("\n" + "=" * 70)
    print("CLEANUP")
    print("=" * 70)
    
    for filename in test_files:
        if os.path.exists(filename):
            os.remove(filename)
            print(f"   🗑️  Đã xóa {filename}")
    
    print("\n✅ Cleanup hoàn thành!")


def main():
    """Chạy tất cả tests."""
    print("=" * 70)
    print("DEMO DOCX LOADER - RAG APPLICATION")
    print("=" * 70)
    
    try:
        # Test 1: Khởi tạo loader
        test_loader_initialization()
        
        # Tạo file test
        if create_test_docx_files():
            # Chạy các test với file đã tạo
            demo_with_created_files()
            
            # Test riêng cho từng feature
            if os.path.exists("test_table.docx"):
                test_table_extraction("test_table.docx")
            
            if os.path.exists("test_formatted.docx"):
                test_formatting_preservation("test_formatted.docx")
            
            if os.path.exists("test_multilang.docx"):
                test_multilingual("test_multilang.docx")
            
            # Test text cleaning
            test_text_cleaning()
        
        print("\n" + "=" * 70)
        print("✅ TẤT CẢ TESTS HOÀN THÀNH")
        print("=" * 70)
        
        print("\n💡 Ví dụ sử dụng:")
        print("─" * 70)
        print("""
from core.loaders.docx_loader import DOCXLoader

# Khởi tạo loader
loader = DOCXLoader(
    enable_table_extraction=True,
    enable_image_extraction=True,
    enable_text_cleaning=True
)

# Load DOCX
docs = loader.load_docx("your_file.docx")

# Xem kết quả
for doc in docs:
    print(doc['text'])
    print(doc['metadata'])
        """)
        print("─" * 70)
        
    except Exception as e:
        print(f"\n❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        # Hỏi user có muốn cleanup không
        response = input("\n🗑️  Xóa các file test? (y/n): ")
        if response.lower() == 'y':
            cleanup_test_files()
        else:
            print("\n📁 Các file test được giữ lại.")


if __name__ == "__main__":
    main()
