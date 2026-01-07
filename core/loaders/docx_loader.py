"""
DOCX Loader nâng cao cho RAG application.

Module này cung cấp các tính năng:
- Trích xuất text từ DOCX thông thường
- Trích xuất và format table từ DOCX
- Xử lý charts/graphs với 3 trường hợp:
  1. Trích xuất caption của đồ thị
  2. Trích xuất bảng số liệu gốc (nếu có)
  3. OCR hoặc mô tả lại nội dung đồ thị
- OCR cho ảnh embedded
- OCR thông minh với bộ lọc (bỏ qua diagrams/charts không cần thiết)
- Làm sạch text cho embedding tối ưu
"""

import os
import io
import unicodedata
from typing import List, Dict, Any
import logging

# Import thư viện DOCX
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

# Import thư viện OCR và xử lý ảnh
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import cv2
import numpy as np

# Cấu hình logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DOCXLoader:
    """DOCX loader với khả năng trích xuất table và OCR ảnh."""
    
    def __init__(
        self,
        ocr_languages: str = "vie+eng+kor+jpn+chi+rus+ara+fra+deu+ita+spa+pol+por+fin+heb+ind",
        enable_image_extraction: bool = True,
        enable_table_extraction: bool = True,
        enable_chart_extraction: bool = True,
        enable_text_cleaning: bool = True,
        min_image_confidence: float = 60.0,
        min_image_words: int = 5
    ):
        """
        Khởi tạo DOCX loader.
        
        Args:
            ocr_languages: Ngôn ngữ cho OCR (format Tesseract)
            enable_image_extraction: Bật/tắt OCR cho ảnh embedded
            enable_table_extraction: Bật/tắt trích xuất table
            enable_chart_extraction: Bật/tắt xử lý đồ thị (caption, bảng gốc, OCR)
            enable_text_cleaning: Bật/tắt làm sạch text
            min_image_confidence: Ngưỡng confidence tối thiểu cho OCR ảnh (0-100)
            min_image_words: Số từ tối thiểu để coi ảnh có text hữu ích
        """
        self.ocr_languages = ocr_languages
        self.enable_image_extraction = enable_image_extraction
        self.enable_table_extraction = enable_table_extraction
        self.enable_chart_extraction = enable_chart_extraction
        self.enable_text_cleaning = enable_text_cleaning
        self.min_image_confidence = min_image_confidence
        self.min_image_words = min_image_words
    
    def load_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Load DOCX và trích xuất text.
        
        Args:
            file_path: Đường dẫn file DOCX
            
        Returns:
            List documents với text và metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Không tìm thấy file DOCX: {file_path}")
        
        doc = Document(file_path)
        
        # Trích xuất tất cả nội dung theo thứ tự
        full_text = self._extract_content_in_order(doc)
        
        # Làm sạch text
        if self.enable_text_cleaning:
            full_text = self._clean_text(full_text)
        
        documents = []
        if full_text.strip():
            documents.append({
                "text": full_text,
                "metadata": {
                    "source": file_path,
                    "num_paragraphs": len(doc.paragraphs),
                    "num_tables": len(doc.tables)
                }
            })
        
        return documents
    
    def _extract_content_in_order(self, doc) -> str:
        """
        Trích xuất nội dung theo thứ tự xuất hiện trong document.
        Bao gồm paragraphs, tables, images, và charts với xử lý đặc biệt.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Text đã trích xuất
        """
        content_parts = []
        
        # Lấy thông tin về tất cả inline shapes (chứa charts) trong document
        chart_contexts = self._extract_chart_contexts(doc) if self.enable_chart_extraction else []
        
        # Iterate qua document body để giữ nguyên thứ tự
        table_counter = 0
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # Paragraph
                paragraph = Paragraph(element, doc)
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            elif isinstance(element, CT_Tbl) and self.enable_table_extraction:
                # Table
                table_counter += 1
                table = Table(element, doc)
                table_text = self._format_table_as_text(table, table_counter)
                if table_text:
                    content_parts.append(table_text)
        
        # Thêm thông tin charts đã trích xuất
        if chart_contexts:
            for chart_info in chart_contexts:
                content_parts.append(chart_info)
        
        # Trích xuất ảnh embedded (nếu bật)
        if self.enable_image_extraction:
            image_text = self._extract_images_from_document(doc)
            if image_text:
                content_parts.append("\n[Text từ ảnh]\n" + image_text)
        
        return "\n\n".join(content_parts)
    
    def _format_table_as_text(self, table: Table, table_idx: int) -> str:
        """
        Format table thành text có cấu trúc dễ đọc (markdown-style).
        
        Args:
            table: python-docx Table object
            table_idx: Index của table trong document
            
        Returns:
            Text được format
        """
        if not table.rows:
            return ""
        
        # Trích xuất data từ table
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            # Bỏ qua hàng rỗng hoàn toàn
            if any(cell for cell in row_data):
                table_data.append(row_data)
        
        if not table_data:
            return ""
        
        # Xác định độ rộng cột
        num_cols = max(len(row) for row in table_data)
        col_widths = [0] * num_cols
        
        for row in table_data:
            for i, cell in enumerate(row):
                if i < num_cols:
                    col_widths[i] = max(col_widths[i], len(cell))
        
        # Format table
        lines = []
        lines.append(f"\n[Bảng {table_idx}]")
        lines.append("=" * 50)
        
        for row_idx, row in enumerate(table_data):
            # Pad các cell để căn chỉnh
            formatted_cells = []
            for i in range(num_cols):
                cell = row[i] if i < len(row) else ""
                formatted_cells.append(cell.ljust(col_widths[i]))
            
            line = " | ".join(formatted_cells)
            lines.append(line)
            
            # Thêm separator sau header (hàng đầu)
            if row_idx == 0 and len(table_data) > 1:
                separator = "-+-".join(["-" * w for w in col_widths])
                lines.append(separator)
        
        lines.append("=" * 50)
        
        logger.info(f"Trích xuất table {table_idx} ({len(table_data)} hàng)")
        return "\n".join(lines)
    
    def _extract_chart_contexts(self, doc) -> List[str]:
        """
        Trích xuất thông tin về charts/graphs với 3 trường hợp:
        1. Có caption (từ paragraph xung quanh)
        2. Có bảng số gốc (source data table)
        3. OCR hoặc mô tả lại chart
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List các đoạn text mô tả chart
        """
        chart_infos = []
        chart_counter = 0
        
        # Tìm tất cả inline shapes (chứa charts) trong document
        for paragraph in doc.paragraphs:
            # Kiểm tra paragraph có chứa inline shape (drawing) không
            if not paragraph._element.xpath('.//w:drawing'):
                continue
            
            chart_counter += 1
            chart_text_parts = [f"\n[Đồ thị {chart_counter}]", "=" * 50]
            
            # TRƯỜNG HỢP 1: Trích xuất caption
            caption = self._extract_chart_caption(paragraph, doc)
            if caption:
                chart_text_parts.append(f"Caption: {caption}")
            
            # TRƯỜNG HỢP 2: Tìm bảng số liệu gốc
            source_table = self._find_chart_source_table(paragraph, doc)
            if source_table:
                chart_text_parts.append(f"Bảng số liệu gốc:\n{source_table}")
            
            # TRƯỜNG HỢP 3: OCR hoặc mô tả chart
            # Lấy ảnh chart và thử OCR (vì một số chart có label, axis text)
            chart_description = self._describe_chart_image(paragraph)
            if chart_description:
                chart_text_parts.append(f"Nội dung đồ thị:\n{chart_description}")
            
            chart_text_parts.append("=" * 50)
            
            # Chỉ thêm nếu có ít nhất một thông tin (caption, data, hoặc description)
            if len(chart_text_parts) > 3:  # Nhiều hơn header và footer
                chart_infos.append("\n".join(chart_text_parts))
                logger.info(f"Trích xuất đồ thị {chart_counter}")
        
        return chart_infos
    
    def _extract_chart_caption(self, chart_paragraph, doc) -> str:
        """
        Trích xuất caption của chart từ paragraph xung quanh.
        
        Heuristic:
        - Tìm paragraph ngay trước/sau chart
        - Caption thường bắt đầu với "Hình", "Biểu đồ", "Figure", "Chart", etc.
        - Hoặc paragraph ngắn (<200 ký tự) gần chart
        
        Args:
            chart_paragraph: Paragraph chứa chart
            doc: Document object
            
        Returns:
            Caption text hoặc empty string
        """
        caption_keywords = ["hình", "biểu đồ", "đồ thị", "figure", "chart", "graph", "diagram"]
        
        # Tìm vị trí của chart paragraph
        try:
            chart_idx = doc.paragraphs.index(chart_paragraph)
        except ValueError:
            return ""
        
        # Kiểm tra paragraph ngay sau (caption thường ở dưới chart)
        if chart_idx + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[chart_idx + 1].text.strip()
            if next_para and len(next_para) < 200:
                # Check có keyword không
                if any(keyword in next_para.lower() for keyword in caption_keywords):
                    return next_para
                # Hoặc paragraph ngắn ngay sau chart cũng có thể là caption
                if len(next_para) < 100:
                    return next_para
        
        # Kiểm tra paragraph ngay trước (caption đôi khi ở trên)
        if chart_idx > 0:
            prev_para = doc.paragraphs[chart_idx - 1].text.strip()
            if prev_para and any(keyword in prev_para.lower() for keyword in caption_keywords):
                return prev_para
        
        return ""
    
    def _find_chart_source_table(self, chart_paragraph, doc) -> str:
        """
        Tìm bảng số liệu gốc tạo ra chart (nếu có).
        
        Heuristic:
        - Bảng gốc thường nằm gần chart (trong 3 elements trước/sau)
        - Bảng có cấu trúc số liệu (nhiều số)
        
        Args:
            chart_paragraph: Paragraph chứa chart
            doc: Document object
            
        Returns:
            Formatted table text hoặc empty string
        """
        # Tìm vị trí chart trong body elements
        chart_element = chart_paragraph._element
        body_elements = list(doc.element.body)
        
        try:
            chart_idx = body_elements.index(chart_element)
        except ValueError:
            return ""
        
        # Tìm trong range 5 elements trước và 3 elements sau
        search_range_before = range(max(0, chart_idx - 5), chart_idx)
        search_range_after = range(chart_idx + 1, min(len(body_elements), chart_idx + 4))
        
        # Ưu tiên tìm sau chart trước (logic: chart → bảng gốc)
        for idx in list(search_range_after) + list(search_range_before):
            element = body_elements[idx]
            if isinstance(element, CT_Tbl):
                table = Table(element, doc)
                # Kiểm tra xem table có phải data table không (chứa nhiều số)
                if self._is_data_table(table):
                    return self._format_table_as_text(table, 0).replace("[Bảng 0]", "").strip()
        
        return ""
    
    def _is_data_table(self, table: Table) -> bool:
        """
        Kiểm tra table có phải là data table (chứa số liệu) không.
        
        Args:
            table: Table object
            
        Returns:
            True nếu là data table
        """
        if not table.rows or len(table.rows) < 2:
            return False
        
        number_count = 0
        total_cells = 0
        
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    total_cells += 1
                    # Kiểm tra cell có chứa số không
                    if any(char.isdigit() for char in text):
                        number_count += 1
        
        # Nếu >= 40% cells chứa số → data table
        if total_cells > 0:
            number_ratio = number_count / total_cells
            return number_ratio >= 0.4
        
        return False
    
    def _describe_chart_image(self, chart_paragraph) -> str:
        """
        OCR hoặc mô tả chart image để trích xuất text labels, axis, values.
        
        Args:
            chart_paragraph: Paragraph chứa chart
            
        Returns:
            Mô tả hoặc OCR text từ chart
        """
        # Thử lấy chart image và OCR
        # Chart trong DOCX thường được embed trong Drawing
        try:
            # Lấy drawing element
            drawings = chart_paragraph._element.xpath('.//w:drawing')
            if not drawings:
                return ""
            
            # Lấy image data từ drawing (nếu chart được render thành image)
            # Note: python-docx không support trực tiếp chart object, 
            # nhưng chart thường có preview image
            
            # Tìm blip (image reference) trong drawing
            blips = drawings[0].xpath('.//a:blip')
            if not blips:
                return "Không thể trích xuất nội dung đồ thị (chưa có ảnh preview)"
            
            # Lấy relationship ID
            rId = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if not rId:
                return ""
            
            # Lấy image part từ relationship
            image_part = chart_paragraph.part.related_parts[rId]
            image_bytes = image_part.blob
            
            # OCR image
            image = Image.open(io.BytesIO(image_bytes))
            
            # Tiền xử lý ảnh
            processed_image = self._preprocess_image(image)
            
            # OCR với config tối ưu cho chart (có axis labels, numbers)
            text = pytesseract.image_to_string(
                processed_image,
                lang=self.ocr_languages,
                config='--psm 11'  # PSM 11: Sparse text (tốt cho chart labels)
            )
            
            text = text.strip()
            if text and len(text) > 10:  # Chỉ return nếu có nội dung đủ
                return f"Text từ OCR: {text}"
            else:
                return "Đồ thị không chứa text có thể đọc được"
            
        except Exception as e:
            logger.debug(f"Không thể OCR chart: {e}")
            return ""
    
    def _extract_images_from_document(self, doc) -> str:
        """
        Trích xuất và OCR text từ ảnh embedded trong DOCX.
        Sử dụng bộ lọc thông minh để tránh OCR các diagram/chart.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Text từ những ảnh có nội dung text hữu ích
        """
        extracted_texts = []
        
        # Lấy tất cả image relationships
        image_parts = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_parts.append(rel.target_part)
        
        if not image_parts:
            return ""
        
        for img_idx, image_part in enumerate(image_parts):
            try:
                # Load ảnh từ binary data
                image_bytes = image_part.blob
                image = Image.open(io.BytesIO(image_bytes))
                
                # Bỏ qua ảnh quá nhỏ (icon/logo)
                if image.width < 50 or image.height < 50:
                    logger.debug(f"Bỏ qua ảnh nhỏ {img_idx} ({image.width}x{image.height})")
                    continue
                
                # Tiền xử lý ảnh
                processed_image = self._preprocess_image(image)
                
                # Lấy kết quả OCR với confidence data
                ocr_data = pytesseract.image_to_data(
                    processed_image,
                    lang=self.ocr_languages,
                    config='--psm 6',
                    output_type=pytesseract.Output.DICT
                )
                
                # Lọc và kiểm tra kết quả OCR
                if not self._is_image_text_meaningful(ocr_data):
                    logger.debug(f"Bỏ qua ảnh {img_idx}: chất lượng thấp hoặc là diagram")
                    continue
                
                # Trích xuất text thực tế
                text = pytesseract.image_to_string(
                    processed_image,
                    lang=self.ocr_languages,
                    config='--psm 6'
                )
                
                text = text.strip()
                if text:
                    extracted_texts.append(text)
                    logger.info(f"Trích xuất text từ ảnh {img_idx} (phát hiện nội dung hữu ích)")
            
            except Exception as e:
                logger.warning(f"Không thể xử lý ảnh {img_idx}: {e}")
                continue
        
        return "\n\n".join(extracted_texts)
    
    def _is_image_text_meaningful(self, ocr_data: Dict) -> bool:
        """
        Xác định ảnh có chứa text hữu ích hay chỉ là diagram/chart.
        
        Args:
            ocr_data: Dữ liệu OCR từ pytesseract.image_to_data
            
        Returns:
            True nếu ảnh có text hữu ích, False nếu là diagram/chart
        """
        # Trích xuất confidence và text
        confidences = []
        words = []
        
        for i, conf in enumerate(ocr_data['conf']):
            # Bỏ qua confidence không hợp lệ
            if conf == -1:
                continue
            
            text = ocr_data['text'][i].strip()
            if text:
                confidences.append(float(conf))
                words.append(text)
        
        # Không tìm thấy text
        if not words:
            return False
        
        # Check 1: Số từ tối thiểu
        if len(words) < self.min_image_words:
            logger.debug(f"Chỉ có {len(words)} từ, dưới ngưỡng {self.min_image_words}")
            return False
        
        # Check 2: Ngưỡng confidence trung bình
        avg_confidence = sum(confidences) / len(confidences)
        if avg_confidence < self.min_image_confidence:
            logger.debug(f"OCR confidence thấp: {avg_confidence:.1f}% < {self.min_image_confidence}%")
            return False
        
        # Check 3: Mật độ text - diagram thường có label ngắn rải rác
        # Tính % "từ thật" (độ dài > 2)
        real_words = [w for w in words if len(w) > 2]
        word_quality_ratio = len(real_words) / len(words) if words else 0
        
        if word_quality_ratio < 0.5:  # Dưới 50% từ thật
            logger.debug(f"Tỉ lệ từ chất lượng thấp: {word_quality_ratio:.2f}")
            return False
        
        # Check 4: Phát hiện chart/diagram
        # Chart thường có nhiều số, ít từ
        number_count = sum(1 for w in words if w.isdigit())
        number_ratio = number_count / len(words) if words else 0
        
        if number_ratio > 0.7:  # Hơn 70% là số → khả năng cao là chart
            logger.debug(f"Tỉ lệ số cao: {number_ratio:.2f} - có thể là chart/diagram")
            return False
        
        logger.debug(f"Ảnh hợp lệ: {len(words)} từ, {avg_confidence:.1f}% conf, {word_quality_ratio:.2f} chất lượng")
        return True
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """
        Tiền xử lý ảnh để tăng độ chính xác OCR.
        
        Args:
            image: PIL Image object
            
        Returns:
            Ảnh đã xử lý
        """
        # Convert sang RGB nếu cần
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Convert sang OpenCV format
        img_array = np.array(image)
        
        # Convert sang grayscale
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        # Khử nhiễu
        denoised = cv2.fastNlMeansDenoising(gray)
        
        # Binarization (Otsu's method)
        _, binary = cv2.threshold(
            denoised, 0, 255,
            cv2.THRESH_BINARY + cv2.THRESH_OTSU
        )
        
        # Convert về PIL
        return Image.fromarray(binary)
    
    def _clean_text(self, text: str) -> str:
        """
        Làm sạch và chuẩn hóa text để tối ưu cho embedding.
        
        Args:
            text: Text thô
            
        Returns:
            Text đã làm sạch
        """
        if not text:
            return text
        
        # 1. Chuẩn hóa Unicode
        text = unicodedata.normalize('NFKC', text)
        
        # 2. Chuẩn hóa khoảng trắng
        text = self._normalize_whitespace(text)
        
        return text.strip()
    
    def _normalize_whitespace(self, text: str) -> str:
        """Chuẩn hóa khoảng trắng thừa nhưng giữ cấu trúc."""
        import re
        
        # Nhiều space thành 1 space
        text = re.sub(r' {2,}', ' ', text)
        
        # Nhiều newline (>2) thành 2 newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Xóa space đầu/cuối mỗi dòng
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text


# Hàm tiện ích để tương thích ngược
def load_docx(
    file_path: str,
    ocr_languages: str = "vie+eng+kor+jpn+chi+rus+ara+fra+deu+ita+spa+pol+por+fin+heb+ind",
    enable_image_extraction: bool = True,
    enable_table_extraction: bool = True,
    enable_chart_extraction: bool = True,
    enable_text_cleaning: bool = True,
    min_image_confidence: float = 60.0,
    min_image_words: int = 5
) -> List[Dict[str, Any]]:
    """
    Load file DOCX và trả về documents với text đã làm sạch.
    
    Args:
        file_path: Đường dẫn file DOCX
        ocr_languages: Ngôn ngữ cho OCR (mặc định: đa ngôn ngữ)
        enable_image_extraction: Bật trích xuất và OCR ảnh embedded
        enable_table_extraction: Bật trích xuất table
        enable_chart_extraction: Bật xử lý đồ thị (caption, bảng gốc, OCR)
        enable_text_cleaning: Bật làm sạch text
        min_image_confidence: Ngưỡng confidence tối thiểu cho OCR ảnh (0-100)
        min_image_words: Số từ tối thiểu để coi ảnh có text hữu ích
        
    Returns:
        List documents với text đã làm sạch và metadata
    """
    loader = DOCXLoader(
        ocr_languages=ocr_languages,
        enable_image_extraction=enable_image_extraction,
        enable_table_extraction=enable_table_extraction,
        enable_chart_extraction=enable_chart_extraction,
        enable_text_cleaning=enable_text_cleaning,
        min_image_confidence=min_image_confidence,
        min_image_words=min_image_words
    )
    return loader.load_docx(file_path)

